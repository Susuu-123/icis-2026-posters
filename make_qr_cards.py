# -*- coding: utf-8 -*-
"""Generate a professional Word document with QR 'scan cards' for both ICIS 2026 posters."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN = RGBColor(0x15, 0x47, 0x34)
ORANGE = RGBColor(0xE8, 0x75, 0x00)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x77, 0x77, 0x77)
FONT = "Arial"

QR_DIR = r"C:\Users\lanli\icis-2026-posters\qr"

CARDS = [
    {
        "qr": QR_DIR + r"\QR_Poster1_SameDiff.png",
        "eyebrow": "UT DALLAS  |  Early Cognition and Communication Lab",
        "heading": "Have a question or comment?",
        "instruction": "Scan to leave a question and we'll follow up by email.",
        "study": "Learning “same” and “different” in infancy",
        "authors": "Elena Luchkina & Elizabeth Spelke",
        "affil": "The University of Texas at Dallas  ·  Harvard University",
    },
    {
        "qr": QR_DIR + r"\QR_Poster2_Unseen.png",
        "eyebrow": "UT DALLAS  |  Early Cognition and Communication Lab",
        "heading": "Interested in this study?",
        "instruction": "Scan to ask a question or share a comment.",
        "study": "How infants represent unseen objects",
        "authors": "Elena Luchkina, Emily Yang & Sandra Waxman",
        "affil": "The University of Texas at Dallas  ·  Northwestern University",
    },
]


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_color)
    tcPr.append(sh)


def set_cell_border(cell, color="154734", sz="18"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def set_cell_margins(cell, top=360, bottom=360, left=360, right=360):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        m.append(el)
    tcPr.append(m)


def para(cell, text=None, size=12, color=DARK, bold=False, italic=False,
         space_before=6, space_after=6, image=None, image_cm=None, caps=False):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if image:
        run = p.add_run()
        run.add_picture(image, width=Cm(image_cm))
        return p
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    if caps:
        rPr = run._element.get_or_add_rPr()
        c = OxmlElement("w:caps")
        c.set(qn("w:val"), "true")
        rPr.append(c)
    return p


def divider(cell, color_hex="E87500", indent_cm=5.0, thickness=28, space_before=2, space_after=10):
    """A short, centered horizontal rule drawn as a real paragraph border (no dash characters)."""
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.left_indent = Cm(indent_cm)
    pf.right_indent = Cm(indent_cm)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pbdr.append(bottom)
    pPr.append(pbdr)
    r = p.add_run(" ")
    r.font.size = Pt(2)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

for i, c in enumerate(CARDS):
    if i > 0:
        doc.add_page_break()
    # small top spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(10)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Cm(14.5)
    set_cell_border(cell, color="154734", sz="20")
    set_cell_margins(cell)
    shade(cell, "FFFFFF")

    # clear the default empty paragraph
    cell.paragraphs[0].text = ""

    para(cell, c["eyebrow"], size=10.5, color=GREEN, bold=True, caps=True,
         space_before=2, space_after=4)
    divider(cell, "E87500")
    para(cell, c["heading"], size=22, color=GREEN, bold=True, space_before=2, space_after=12)
    para(cell, image=c["qr"], image_cm=6.6, space_before=4, space_after=12)
    para(cell, c["instruction"], size=13.5, color=DARK, space_before=2, space_after=14)
    para(cell, c["study"], size=11.5, color=GRAY, italic=True, space_before=2, space_after=10)
    # bottom green rule
    divider(cell, "154734", thickness=18, space_before=4, space_after=8)
    para(cell, c["authors"], size=11.5, color=GREEN, bold=True, space_before=2, space_after=1)
    para(cell, c["affil"], size=9.5, color=GRAY, space_before=0, space_after=2)

out = r"C:\Users\lanli\icis-2026-posters\ICIS2026_QR_Scan_Cards.docx"
doc.save(out)
print("SAVED:", out)
